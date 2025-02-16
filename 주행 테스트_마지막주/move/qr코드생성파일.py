from io import BytesIO
import json
import qrcode
from docx import Document
from docx.shared import Inches

# Word 문서 생성
document = Document()

# 9행×9열 표 생성
num_rows = 9
num_cols = 9
table = document.add_table(rows=num_rows, cols=num_cols)

# 각 셀에 들어갈 QR 코드와 좌표 텍스트 삽입
# Word 표의 첫 번째 행은 문서 상단에 위치하므로,
# y 좌표는 8부터 0까지 (top -> bottom)로 매핑합니다.
for row in range(num_rows):
    y_coord = 8 - row  # top row: y=8, bottom row: y=0
    for col in range(num_cols):
        x_coord = col  # x: 0 ~ 8
        
        # QR 코드에 저장할 JSON 데이터 생성
        data = {"location_y": y_coord, "location_x": x_coord}
        data_str = json.dumps(data)
        
        # QR 코드 생성 (box_size=5로 크기를 5배로 설정)
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=5,
            border=1,
        )
        qr.add_data(data_str)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        
        # QR 이미지를 메모리 내 BytesIO 객체에 저장 (PNG 형식)
        bio = BytesIO()
        img.save(bio, format="PNG")
        bio.seek(0)
        
        # Word 표의 해당 셀에 이미지와 텍스트 삽입
        cell = table.cell(row, col)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        # QR 코드 이미지 삽입 (너비 1인치로 조정, 필요에 따라 조절)
        run.add_picture(bio, width=Inches(1))
        # 이미지 아래에 좌표 텍스트 추가 (새 줄로)
        paragraph.add_run(f"\n({y_coord},{x_coord})")

# Word 파일 저장
output_filename = "qr_codes.docx"
document.save(output_filename)
print(f"Word 파일 '{output_filename}'이(가) 생성되었습니다.")
